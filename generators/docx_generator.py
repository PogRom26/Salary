from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def money(value):

    try:
        return f"{float(value):,.2f}".replace(",", " ")
    except Exception:
        return "0.00"

def set_cell_color(cell, color):

    cell._tc.get_or_add_tcPr().append(
        parse_xml(
            rf'<w:shd {nsdecls("w")} '
            rf'w:fill="{color}"/>'
        )
    )


def add_table(doc, headers, rows):

    table = doc.add_table(
        rows=1,
        cols=len(headers)
    )

    table.style = "Table Grid"

    hdr = table.rows[0].cells

    for i, header in enumerate(headers):

        hdr[i].text = str(header)

        set_cell_color(
            hdr[i],
            "D9EAF7"
        )

    for row_data in rows:

        row = table.add_row().cells

        for i, value in enumerate(row_data):

            row[i].text = str(value)

        first_col = str(row_data[0])

        # Серый для ИТОГО

        if "ИТОГО" in first_col:

            for cell in row:
                set_cell_color(
                    cell,
                    "F2F2F2"
                )

        # Зеленый для зарплаты

        elif (
                "Размер заработной платы"
                in first_col
        ):

            for cell in row:
                set_cell_color(
                    cell,
                    "DFF0D8"
                )

    doc.add_paragraph("")


def generate_docx(
    docx_path,
    employee,
    report,
    year,
    month,
):

    profit = report["profit"]
    brand = report["brand"]
    debt = report["debt"]
    cycle = report["cycle"]
    communications = report["communications"]
    timesheet = report["timesheet"]

    doc = Document()

    MONTHS = {
        1: "январь",
        2: "февраль",
        3: "март",
        4: "апрель",
        5: "май",
        6: "июнь",
        7: "июль",
        8: "август",
        9: "сентябрь",
        10: "октябрь",
        11: "ноябрь",
        12: "декабрь",
    }

    style = doc.styles["Normal"]

    style.font.name = "DejaVu Sans"
    style.font.size = Pt(10)

    # ==========================================
    # Заголовок
    # ==========================================

    title = doc.add_heading(
        f"Расчёт сотрудника b2b-направления "
        f"за {MONTHS[month]} {year} года",
        level=1
    )

    for run in title.runs:
        run.font.name = "DejaVu Sans"
        run.bold = True

    doc.add_paragraph("")

    employee_title = doc.add_heading(
        employee,
        level=2
    )

    for run in employee_title.runs:
        run.font.name = "DejaVu Sans"
        run.bold = True

    doc.add_paragraph("")

    for run in title.runs:
        run.font.name = "DejaVu Sans"
        run.font.size = Pt(16)
        run.bold = True

    # ==========================================
    # Доход
    # ==========================================

    doc.add_heading(
        "1. Доход",
        level=2
    )

    add_table(
        doc,
        ["Показатель", "Значение"],
        [
            ["Сумма продаж", money(profit["sales"])],
            ["Доход текущий", money(profit["income"])],
            ["Рентабельность текущая",
             f"{profit['profitability']:.2f}%"],
            ["Бонус от дохода 5%",
             money(profit["bonus"])],
        ]
    )

    # ==========================================
    # KPI ZIC
    # ==========================================

    doc.add_heading(
        "2. KPI ZIC",
        level=2
    )

    rows = []

    for item in brand["zic"]:

        rows.append([
            item["brand"],
            money(item["plan"]),
            money(item["fact"]),
            f"{item['percent']:.2f}%",
            money(item["bonus"]),
        ])

    add_table(
        doc,
        [
            "Показатель",
            "План",
            "Факт",
            "% выполнения",
            "Бонус",
        ],
        rows
    )

    # ==========================================
    # Лукойл
    # ==========================================

    doc.add_heading(
        "3. Лукойл",
        level=2
    )

    if brand["lukoil"]:

        add_table(
            doc,
            ["Показатель", "Значение"],
            [
                [
                    "Объем продаж, кг",
                    money(
                        brand["lukoil"]["kg"]
                    ),
                ],
                [
                    "Ставка",
                    money(
                        brand["lukoil"]["rate"]
                    ),
                ],
                [
                    "Бонус",
                    money(
                        brand["lukoil"]["bonus"]
                    ),
                ],
            ]
        )

    # ==========================================
    # Другие KPI
    # ==========================================

    doc.add_heading(
        "4. Другие KPI",
        level=2
    )

    rows = []

    total_percent = 0

    for item in brand["other"]:

        total_percent += (
            item["weighted_percent"]
        )

        rows.append([
            item["brand"],
            f"{item['percent']:.2f}%",
            f"{item['weight']:.2f}%",
            f"{item['weighted_percent']:.2f}%",
        ])

    rows.append([
        "ИТОГО",
        "",
        "",
        f"{total_percent:.2f}%"
    ])

    rows.append([
        "Бонус",
        "",
        "",
        money(
            brand["other_bonus_total"]
        )
    ])

    add_table(
        doc,
        [
            "Показатель",
            "% выполнения",
            "Вес",
            "Итог %",
        ],
        rows
    )

    # ==========================================
    # Дебиторка
    # ==========================================

    doc.add_heading(
        "5. Дебиторская задолженность",
        level=2
    )

    add_table(
        doc,
        ["Показатель", "Значение"],
        [
            [
                "Общая просроченная задолженность",
                money(debt["total"]),
            ],

            [
                "Ответственность за ПДЗ",
                "-" + money(
                    debt["responsibility"]
                ),
            ],

            [
                "Отношение ПДЗ к обороту",
                f"{debt['ratio']:.2f}%"
            ],
        ]
    )

    # ==========================================
    # Цикл сделки
    # ==========================================

    doc.add_heading(
        "6. Цикл сделки",
        level=2
    )

    add_table(
        doc,
        ["Показатель", "Значение"],
        [
            [
                "Цикл сделки, план",
                f"{cycle['plan']:.1f}"
            ],

            [
                "Цикл сделки, факт",
                f"{cycle['fact']:.1f}"
            ],

            [
                "Соотношение",
                f"{cycle['ratio'] * 100:.2f}%"
            ],

            [
                "Бонус",
                money(
                    cycle["bonus"]
                )
            ],
        ]
    )

    # ==========================================
    # Коммуникации
    # ==========================================

    doc.add_heading(
        "7. Коммуникации за месяц",
        level=2
    )

    rows = []

    for name, data in communications.items():

        rows.append([
            name.capitalize(),
            data["total"],
            data["success"],
            data["failed"],
        ])

    add_table(
        doc,
        [
            "Тип",
            "Всего",
            "Успешно",
            "Неудачно",
        ],
        rows
    )

    # ==========================================
    # Итог
    # ==========================================

    total_bonus = (
        profit["bonus"]
        + brand["zic_bonus_total"]
        + brand["lukoil_bonus"]
        + brand["other_bonus_total"]
        + cycle["bonus"]
        - debt["responsibility"]
    )

    salary_total = (
        total_bonus
        + timesheet["salary"]
    )

    doc.add_heading(
        "Итоговый расчёт",
        level=2
    )

    add_table(
        doc,
        ["Показатель", "Сумма"],
        [
            [
                "Бонус от дохода 5%",
                money(
                    profit["bonus"]
                ),
            ],

            [
                "Бонус KPI ZIC",
                money(
                    brand["zic_bonus_total"]
                ),
            ],

            [
                "Бонус Лукойл",
                money(
                    brand["lukoil_bonus"]
                ),
            ],

            [
                "Бонус за другие KPI",
                money(
                    brand["other_bonus_total"]
                ),
            ],

            [
                "Бонус за цикл сделки",
                money(
                    cycle["bonus"]
                ),
            ],

            [
                "Ответственность за ПДЗ",
                "-" + money(
                    debt["responsibility"]
                ),
            ],

            [
                "ИТОГО, премия",
                money(total_bonus),
            ],

            [
                "Отработано часов",
                f"{timesheet['hours']:.2f}",
            ],

            [
                "Оклад",
                money(
                    timesheet["salary"]
                ),
            ],

            [
                "Размер заработной платы",
                money(
                    salary_total
                ),
            ],
        ]
    )

    doc.save(docx_path)